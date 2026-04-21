"""
Drug Classification Document Generator (BNF Format)
====================================================
Generates a comprehensive Word document with drug classifications
organized by body system and sub-topic.

Usage:
    python create_drug_document.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configuration
OUTPUT_FILE = "Drug_Classification_2_BNF.docx"
FONT_NAME = "Times New Roman"
FONT_SIZE = 11
DRUG_HEADING_SIZE = 14
SYSTEM_HEADING_SIZE = 16
SUBTOPIC_HEADING_SIZE = 13


def set_font(run, name=FONT_NAME, size=FONT_SIZE, bold=False, italic=False):
    """Set font properties for a run."""
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        r.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)


def add_system_heading(doc, text, page_break=True):
    """Add a body system heading (e.g., 'Cardiovascular System')."""
    para = doc.add_paragraph()
    if page_break:
        run_br = para.add_run()
        run_br.add_break(WD_BREAK.PAGE)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text.upper())
    set_font(run, size=SYSTEM_HEADING_SIZE, bold=True)
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.space_before = Pt(24)


def add_subtopic_heading(doc, number, text):
    """Add a sub-topic heading (e.g., '1. Anticoagulants')."""
    para = doc.add_paragraph()
    run = para.add_run(f"{number}. {text}")
    set_font(run, size=SUBTOPIC_HEADING_SIZE, bold=True)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.space_before = Pt(12)


def add_drug_name_heading(doc, text):
    """Add the drug name as a bold heading on a new page."""
    para = doc.add_paragraph()
    # Page break before each drug
    para.paragraph_format.page_break_before = True
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(text)
    set_font(run, size=DRUG_HEADING_SIZE, bold=True)
    para.paragraph_format.space_after = Pt(6)


def add_photo_placeholder(doc, description):
    """Add a photo placeholder in the required format."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"[ Photo: {description} \u2013 4 cm \u00d7 4 cm ]")
    set_font(run, italic=True)
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(8)


def add_labeled_line(doc, label, text):
    """Add 'Label: text' on a single line."""
    para = doc.add_paragraph()
    run_label = para.add_run(f"{label}: ")
    set_font(run_label, bold=True)
    run_text = para.add_run(text)
    set_font(run_text)
    para.paragraph_format.space_after = Pt(3)
    para.paragraph_format.space_before = Pt(2)


def add_labeled_paragraph(doc, label, text):
    """Add a bold label followed by a paragraph of text."""
    para = doc.add_paragraph()
    run_label = para.add_run(f"{label}: ")
    set_font(run_label, bold=True)
    run_text = para.add_run(text)
    set_font(run_text)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.space_before = Pt(4)


def add_section_label(doc, label):
    """Add a bold section label (e.g., 'Indications:')."""
    para = doc.add_paragraph()
    run = para.add_run(f"{label}:")
    set_font(run, bold=True)
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.space_before = Pt(4)


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point. Optionally bold a prefix portion."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(2)
    para.paragraph_format.left_indent = Cm(1.27)
    # Add bullet character
    if bold_prefix:
        run_b = para.add_run(f"\u2022 {bold_prefix}: ")
        set_font(run_b, bold=True)
        run_t = para.add_run(text)
        set_font(run_t)
    else:
        run = para.add_run(f"\u2022 {text}")
        set_font(run)


def add_drug_entry(doc, drug):
    """Add a complete drug entry to the document."""
    d = drug

    # Drug name heading (new page)
    add_drug_name_heading(doc, d["name"])

    # Photo placeholder
    add_photo_placeholder(doc, d["photo_desc"])

    # Generic Name
    add_labeled_line(doc, "Generic Name", d["generic_name"])

    # Trade Name
    add_labeled_line(doc, "Trade Name", d["trade_name"])

    # Mechanism of Action (paragraph)
    add_labeled_paragraph(doc, "Mechanism of Action", d["mechanism"])

    # Indications (bullets)
    add_section_label(doc, "Indications")
    for item in d["indications"]:
        add_bullet(doc, item)

    # Dosage (bullets)
    add_section_label(doc, "Dosage")
    for item in d["dosage"]:
        add_bullet(doc, item)

    # Route of Administration
    add_labeled_line(doc, "Route of Administration", d["route"])

    # Pharmaceutical Preparation
    add_labeled_line(doc, "Pharmaceutical Preparation", d["preparation"])

    # Strength
    add_labeled_line(doc, "Strength", d["strength"])

    # Contraindications (bullets)
    add_section_label(doc, "Contraindications")
    for item in d["contraindications"]:
        add_bullet(doc, item)

    # Side Effects (bullets with bold prefix)
    add_section_label(doc, "Side Effects")
    for item in d["side_effects"]:
        if ":" in item:
            prefix, rest = item.split(":", 1)
            add_bullet(doc, rest.strip(), bold_prefix=prefix.strip())
        else:
            add_bullet(doc, item)

    # Nursing Considerations (bullets with bold prefix)
    add_section_label(doc, "Nursing Considerations")
    for item in d["nursing_considerations"]:
        if ":" in item:
            prefix, rest = item.split(":", 1)
            add_bullet(doc, rest.strip(), bold_prefix=prefix.strip())
        else:
            add_bullet(doc, item)

    # Storage Guidelines
    add_labeled_paragraph(doc, "Storage Guidelines", d["storage"])


def build_document(all_data):
    """Build the complete document from structured data."""
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE)
    font.color.rgb = RGBColor(0, 0, 0)

    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # Title Page
    for _ in range(6):
        doc.add_paragraph()
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("DRUG CLASSIFICATION")
    set_font(run, size=22, bold=True)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_para.add_run("Based on British National Formulary (BNF) Guidelines")
    set_font(run, size=14)

    subtitle2 = doc.add_paragraph()
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle2.add_run("Cardiovascular System & Urinary System")
    set_font(run, size=13, italic=True)

    # Process each system
    first_system = True
    for system_name, subtopics in all_data.items():
        add_system_heading(doc, system_name, page_break=True)

        for sub_num, (subtopic_name, drugs) in enumerate(subtopics.items(), 1):
            add_subtopic_heading(doc, sub_num, subtopic_name)

            for drug in drugs:
                add_drug_entry(doc, drug)

    doc.save(OUTPUT_FILE)
    print(f"\nDocument saved: {OUTPUT_FILE}")
    print(f"Total drugs: {sum(len(drugs) for st in all_data.values() for drugs in st.values())}")


if __name__ == "__main__":
    from drug_data_cardio import CARDIOVASCULAR_DATA
    from drug_data_urinary import URINARY_DATA

    all_data = {
        "Cardiovascular System": CARDIOVASCULAR_DATA,
        "Urinary System": URINARY_DATA,
    }
    build_document(all_data)
    print("Done!")
