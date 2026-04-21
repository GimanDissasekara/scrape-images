"""
Medicine Photo Scraper & Document Inserter
==========================================
This script:
1. Reads the Drug_Classification_BNF.docx document
2. Extracts all medicine names from "[ Photo: <name> – 4 cm × 4 cm ]" placeholders
3. Downloads photos of each medicine from Bing Images
4. Saves photos to a 'medicine_photos' folder
5. Inserts the downloaded photos into the document, replacing the placeholder text

Usage:
    python medicine_photo_scraper.py
"""

import os
import re
import sys
import time
import json
import logging
import hashlib
import urllib.parse
import urllib.request
from pathlib import Path

from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Configuration ──────────────────────────────────────────────────────────
DOCX_FILE = "Drug_Classification_BNF.docx"
OUTPUT_DOCX = "Drug_Classification_BNF_with_photos.docx"
PHOTO_DIR = "medicine_photos"
IMAGE_SIZE_CM = 4  # 4 cm × 4 cm as specified in the document
MAX_IMAGES_PER_MEDICINE = 1

# Regex pattern to match photo placeholders
PHOTO_PATTERN = re.compile(
    r'\[\s*Photo space:\s*(.+?)\s*\]'
)

# ── Logging Setup ──────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s [%(levelname)s] %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('scraper.log', encoding='utf-8')
    ]
)
logger = logging.getLogger(__name__)

# User-Agent for HTTP requests
USER_AGENT = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/120.0.0.0 Safari/537.36"
)


def extract_medicine_names(doc_path: str) -> list[dict]:
    """
    Extract medicine names and their paragraph indices from the document.
    Returns: [{'name': str, 'para_index': int, 'full_match': str, 'full_text': str}, ...]
    """
    doc = Document(doc_path)
    medicines = []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = PHOTO_PATTERN.search(text)
        if match:
            medicine_name = match.group(1).strip()
            medicines.append({
                'name': medicine_name,
                'para_index': i,
                'full_match': match.group(0),
                'full_text': text,
            })

    return medicines


def create_safe_filename(name: str) -> str:
    """Convert a medicine name to a safe filename."""
    safe = re.sub(r'[^\w\s-]', '', name)
    safe = re.sub(r'\s+', '_', safe)
    return safe.lower()[:80]


def get_existing_image(med_dir: str) -> str | None:
    """Check if an image already exists in the directory."""
    if os.path.exists(med_dir):
        for f in os.listdir(med_dir):
            if f.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp')):
                return os.path.join(med_dir, f)
    return None


def download_image_from_url(url: str, save_path: str) -> bool:
    """Download a single image from URL."""
    try:
        req = urllib.request.Request(url, headers={'User-Agent': USER_AGENT})
        with urllib.request.urlopen(req, timeout=15) as response:
            content_type = response.headers.get('Content-Type', '')
            if 'image' not in content_type and not url.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.webp')):
                return False
            data = response.read()
            if len(data) < 1000:  # Skip tiny images (likely broken)
                return False
            with open(save_path, 'wb') as f:
                f.write(data)
            return True
    except Exception:
        return False


def search_bing_images(query: str, count: int = 5) -> list[str]:
    """
    Search Bing Images and return image URLs using web scraping.
    """
    encoded_query = urllib.parse.quote(query)
    url = f"https://www.bing.com/images/search?q={encoded_query}&form=HDRSC2&first=1"

    try:
        req = urllib.request.Request(url, headers={
            'User-Agent': USER_AGENT,
            'Accept': 'text/html,application/xhtml+xml',
            'Accept-Language': 'en-US,en;q=0.9',
        })
        with urllib.request.urlopen(req, timeout=15) as response:
            html = response.read().decode('utf-8', errors='ignore')
        
        # Extract image URLs from Bing's HTML
        # Bing stores image data in 'murl' parameter
        image_urls = []
        
        # Method 1: Extract from murl in data attributes
        murl_pattern = re.compile(r'"murl"\s*:\s*"(https?://[^"]+)"')
        matches = murl_pattern.findall(html)
        image_urls.extend(matches[:count])
        
        # Method 2: Extract from img src attributes if needed
        if not image_urls:
            img_pattern = re.compile(r'<img[^>]+src="(https?://[^"]+)"[^>]*>')
            matches = img_pattern.findall(html)
            for m in matches:
                if 'bing.com' not in m and 'microsoft.com' not in m:
                    image_urls.append(m)
                    if len(image_urls) >= count:
                        break

        return image_urls[:count]

    except Exception as e:
        logger.debug(f"Bing search failed for '{query}': {e}")
        return []


def search_duckduckgo_images(query: str, count: int = 5) -> list[str]:
    """
    Search DuckDuckGo Images as a fallback.
    """
    try:
        # First get the vqd token
        encoded_query = urllib.parse.quote(query)
        token_url = f"https://duckduckgo.com/?q={encoded_query}&iax=images&ia=images"
        
        req = urllib.request.Request(token_url, headers={
            'User-Agent': USER_AGENT,
        })
        with urllib.request.urlopen(req, timeout=15) as response:
            html = response.read().decode('utf-8', errors='ignore')
        
        # Extract vqd token
        vqd_match = re.search(r'vqd=["\']([^"\']+)', html)
        if not vqd_match:
            return []
        
        vqd = vqd_match.group(1)
        
        # Search images API
        api_url = (
            f"https://duckduckgo.com/i.js?l=us-en&o=json&q={encoded_query}"
            f"&vqd={vqd}&f=,,,,,&p=1"
        )
        
        req = urllib.request.Request(api_url, headers={
            'User-Agent': USER_AGENT,
            'Referer': 'https://duckduckgo.com/',
        })
        with urllib.request.urlopen(req, timeout=15) as response:
            data = json.loads(response.read().decode('utf-8'))
        
        image_urls = []
        for result in data.get('results', [])[:count]:
            if 'image' in result:
                image_urls.append(result['image'])
        
        return image_urls

    except Exception as e:
        logger.debug(f"DuckDuckGo search failed for '{query}': {e}")
        return []


def download_medicine_photos(medicines: list[dict], photo_dir: str) -> dict[str, str]:
    """
    Download photos for each unique medicine.
    Uses Bing scraping, with DuckDuckGo as fallback.
    Returns: {medicine_name: local_image_path}
    """
    os.makedirs(photo_dir, exist_ok=True)
    image_paths = {}

    # Get unique medicine names
    unique_medicines = {}
    for med in medicines:
        if med['name'] not in unique_medicines:
            unique_medicines[med['name']] = med

    unique_count = len(unique_medicines)
    logger.info(f"Found {len(medicines)} placeholders, {unique_count} unique medicines")

    for idx, (name, med) in enumerate(unique_medicines.items(), 1):
        safe_name = create_safe_filename(name)
        med_dir = os.path.join(photo_dir, safe_name)
        os.makedirs(med_dir, exist_ok=True)

        # Check if already downloaded
        existing = get_existing_image(med_dir)
        if existing:
            image_paths[name] = existing
            logger.info(f"[{idx}/{unique_count}] Already exists: {name}")
            continue

        # Build search query
        search_query = f"{name} medicine tablet packaging"
        logger.info(f"[{idx}/{unique_count}] Searching: {name}")

        # Try Bing first
        image_urls = search_bing_images(search_query)
        
        # Fallback to DuckDuckGo
        if not image_urls:
            logger.info(f"  Trying DuckDuckGo fallback...")
            image_urls = search_duckduckgo_images(search_query)

        # Fallback: Try simpler query
        if not image_urls:
            simpler_query = name.split('/')[0].strip() + " medicine"
            image_urls = search_bing_images(simpler_query)

        # Download the first working image
        downloaded = False
        for img_url in image_urls:
            ext = '.jpg'
            for e in ['.png', '.jpeg', '.gif', '.webp']:
                if e in img_url.lower():
                    ext = e
                    break

            save_path = os.path.join(med_dir, f"photo{ext}")
            if download_image_from_url(img_url, save_path):
                image_paths[name] = save_path
                logger.info(f"  ✓ Downloaded successfully")
                downloaded = True
                break
            else:
                # Clean up failed download
                if os.path.exists(save_path):
                    os.remove(save_path)

        if not downloaded:
            logger.warning(f"  ✗ No image found for: {name}")

        # Polite delay between searches
        time.sleep(1.5)

    return image_paths


def download_with_icrawler_bing(medicines: list[dict], photo_dir: str, 
                                 existing_paths: dict[str, str]) -> dict[str, str]:
    """
    Try icrawler's BingImageCrawler for any remaining missing images.
    """
    try:
        from icrawler.builtin import BingImageCrawler
    except ImportError:
        logger.info("icrawler not available, skipping Bing crawler fallback")
        return {}

    image_paths = {}
    unique_medicines = {}
    for med in medicines:
        if med['name'] not in unique_medicines and med['name'] not in existing_paths:
            unique_medicines[med['name']] = med

    if not unique_medicines:
        return {}

    logger.info(f"\n[icrawler Fallback] Trying Bing crawler for {len(unique_medicines)} missing medicines")

    for idx, (name, med) in enumerate(unique_medicines.items(), 1):
        safe_name = create_safe_filename(name)
        med_dir = os.path.join(photo_dir, safe_name)
        os.makedirs(med_dir, exist_ok=True)

        existing = get_existing_image(med_dir)
        if existing:
            image_paths[name] = existing
            continue

        search_query = f"{name} medicine"
        logger.info(f"  [{idx}/{len(unique_medicines)}] Crawling: {name}")

        try:
            crawler = BingImageCrawler(
                storage={'root_dir': med_dir},
                log_level=logging.ERROR
            )
            crawler.crawl(keyword=search_query, max_num=1)

            existing = get_existing_image(med_dir)
            if existing:
                image_paths[name] = existing
                logger.info(f"    ✓ Downloaded via icrawler")
        except Exception as e:
            logger.debug(f"    ✗ icrawler error: {e}")

        time.sleep(1)

    return image_paths


def insert_photos_into_document(doc_path: str, output_path: str,
                                 medicines: list[dict], image_paths: dict[str, str]):
    """
    Insert downloaded photos into the document, replacing the photo placeholder
    paragraphs with the actual images.
    """
    doc = Document(doc_path)
    
    inserted = 0
    missing = 0

    # Process in reverse order so paragraph indices remain valid
    for med in reversed(medicines):
        name = med['name']
        para_idx = med['para_index']
        para = doc.paragraphs[para_idx]

        if name in image_paths and os.path.exists(image_paths[name]):
            img_path = image_paths[name]

            # Clear the placeholder text
            para.clear()

            # Add the image
            run = para.add_run()
            try:
                run.add_picture(img_path, width=Cm(IMAGE_SIZE_CM), height=Cm(IMAGE_SIZE_CM))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                inserted += 1
                logger.info(f"  ✓ Inserted photo for: {name}")
            except Exception as e:
                # If image is corrupted or unsupported, restore placeholder
                para.add_run(med['full_text'])
                logger.error(f"  ✗ Failed to insert photo for {name}: {e}")
                missing += 1
        else:
            logger.warning(f"  ✗ No photo available for: {name}")
            missing += 1

    doc.save(output_path)
    logger.info(f"\nDocument saved: {output_path}")
    logger.info(f"Photos inserted: {inserted}, Missing: {missing}")

    return inserted, missing


def print_summary(medicines, image_paths, inserted, missing):
    """Print a summary of the operation."""
    print("\n" + "=" * 60)
    print("   MEDICINE PHOTO SCRAPER - SUMMARY")
    print("=" * 60)
    print(f"  Total photo placeholders found : {len(medicines)}")
    
    unique_names = set(m['name'] for m in medicines)
    print(f"  Unique medicines               : {len(unique_names)}")
    print(f"  Photos downloaded              : {len(image_paths)}")
    print(f"  Photos inserted into document  : {inserted}")
    print(f"  Missing/failed                 : {missing}")
    print(f"  Output document                : {OUTPUT_DOCX}")
    print(f"  Photos saved in                : {PHOTO_DIR}/")
    print("=" * 60)

    if missing > 0:
        print("\n  Medicines without photos:")
        for name in sorted(unique_names):
            if name not in image_paths:
                print(f"    - {name}")
        print()


def main():
    logger.info("=" * 60)
    logger.info("  Medicine Photo Scraper & Document Inserter")
    logger.info("=" * 60)

    # Check source document exists
    if not os.path.exists(DOCX_FILE):
        logger.error(f"Document not found: {DOCX_FILE}")
        sys.exit(1)

    # Step 1: Extract medicine names from document
    logger.info("\nStep 1: Extracting medicine names from document...")
    medicines = extract_medicine_names(DOCX_FILE)
    logger.info(f"   Found {len(medicines)} photo placeholders")

    if not medicines:
        logger.error("No photo placeholders found in the document!")
        sys.exit(1)

    # Print all found medicines
    unique_names = sorted(set(m['name'] for m in medicines))
    logger.info(f"\n   Unique medicines ({len(unique_names)}):")
    for i, name in enumerate(unique_names, 1):
        logger.info(f"   {i:3d}. {name}")

    # Step 2: Download photos
    logger.info("\nStep 2: Downloading medicine photos...")
    image_paths = download_medicine_photos(medicines, PHOTO_DIR)

    # Step 2b: Try icrawler for missing images
    missing_names = [n for n in unique_names if n not in image_paths]
    if missing_names:
        logger.info(f"\nStep 2b: {len(missing_names)} medicines still missing, trying icrawler...")
        extra_paths = download_with_icrawler_bing(medicines, PHOTO_DIR, image_paths)
        image_paths.update(extra_paths)

    logger.info(f"\n   Total photos downloaded: {len(image_paths)}/{len(unique_names)}")

    # Step 3: Insert photos into document
    logger.info(f"\nStep 3: Inserting photos into document...")
    inserted, missing = insert_photos_into_document(
        DOCX_FILE, OUTPUT_DOCX, medicines, image_paths
    )

    # Print summary
    print_summary(medicines, image_paths, inserted, missing)

    logger.info("Done!")


if __name__ == "__main__":
    main()
